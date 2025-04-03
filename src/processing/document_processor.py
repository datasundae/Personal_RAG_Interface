"""Document processing module for the RAG system."""

import hashlib
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
import uuid

import numpy as np
from sentence_transformers import SentenceTransformer
from PyPDF2 import PdfReader
import pytesseract
from pdf2image import convert_from_path

from ..config.config import DOC_CONFIG, MODEL_CONFIG
from ..database.sqlite_vector_db import SQLiteVectorDB

from .rag_document import RAGDocument
from .pdf_processor import PDFProcessor
from .image_processor import ImageProcessor

from typing import List, Optional, Dict, Any, Tuple
import PyPDF2
import io
import os
from pathlib import Path
import pytesseract
from pdf2image import convert_from_path
import cv2
import numpy as np
from src.processing.rag_document import RAGDocument
from src.processing.pdf_processor import PDFProcessor
from src.processing.image_processor import ImageProcessor
import json
import csv
import pandas as pd
import docx
from PIL import Image
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse

class DocumentProcessor:
    """Process different types of documents and convert them to RAGDocument objects."""
    
    def __init__(self):
        """Initialize the document processor."""
        self.pdf_processor = PDFProcessor()
        self.image_processor = ImageProcessor()
    
    def process_document(self, path: str, metadata: Optional[Dict[str, Any]] = None) -> RAGDocument:
        """Process a document and return a RAGDocument.
        
        Args:
            path: Path to the document or URL
            metadata: Optional metadata to attach to the document
            
        Returns:
            RAGDocument containing the text content and metadata
        """
        try:
            # Check if path is a URL
            if self._is_url(path):
                return self._process_url(path, metadata)
            
            # Get file extension
            _, ext = os.path.splitext(path)
            ext = ext.lower()
            
            # Process based on file type
            if ext == '.pdf':
                return self.pdf_processor.process_pdf(path, metadata)
            elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self.image_processor.process_image(path, metadata)
            elif ext == '.docx':
                return self._process_docx(path, metadata)
            elif ext == '.txt':
                return self._process_text(path, metadata)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
            
        except Exception as e:
            raise ValueError(f"Error processing document {path}: {str(e)}")
    
    def process_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[RAGDocument]:
        """
        Process a single file and convert it to RAGDocument objects.
        
        Args:
            file_path: Path to the file to process
            metadata: Optional metadata to add to the documents
            
        Returns:
            List of RAGDocument objects
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Initialize metadata if None
        if metadata is None:
            metadata = {}
        
        # Add file-specific metadata
        file_metadata = {
            "source": str(file_path),
            "file_type": file_path.suffix.lower(),
            "file_name": file_path.name,
            **metadata
        }
        
        # Process based on file type
        if file_path.suffix.lower() == '.pdf':
            return self.pdf_processor.process_pdf(file_path, file_metadata)
        elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return self.image_processor.process_image(file_path, file_metadata)
        elif file_path.suffix.lower() == '.txt':
            return self._process_text_file(file_path, file_metadata)
        elif file_path.suffix.lower() == '.json':
            return self._process_json_file(file_path, file_metadata)
        elif file_path.suffix.lower() == '.csv':
            return self._process_csv_file(file_path, file_metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
    
    def process_directory(self, directory_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[RAGDocument]:
        """
        Process all supported files in a directory.
        
        Args:
            directory_path: Path to the directory to process
            metadata: Optional metadata to add to all documents
            
        Returns:
            List of RAGDocument objects
        """
        directory_path = Path(directory_path)
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        documents = []
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and self._is_supported_file(file_path):
                try:
                    file_documents = self.process_file(str(file_path), metadata)
                    documents.extend(file_documents)
                except Exception as e:
                    print(f"Error processing {file_path}: {str(e)}")
        
        return documents
    
    def _is_supported_file(self, file_path: Path) -> bool:
        """Check if the file type is supported."""
        supported_extensions = {'.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.txt', '.json', '.csv'}
        return file_path.suffix.lower() in supported_extensions
    
    def _process_text_file(self, file_path: Path, metadata: Dict[str, Any]) -> List[RAGDocument]:
        """Process a text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return [RAGDocument(content=content, metadata=metadata)]
    
    def _process_json_file(self, file_path: Path, metadata: Dict[str, Any]) -> List[RAGDocument]:
        """Process a JSON file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        documents = []
        if isinstance(data, dict):
            # If JSON is a single object, convert to string
            content = json.dumps(data, ensure_ascii=False)
            documents.append(RAGDocument(content=content, metadata=metadata))
        elif isinstance(data, list):
            # If JSON is an array, process each item
            for i, item in enumerate(data):
                item_metadata = {**metadata, "item_index": i}
                content = json.dumps(item, ensure_ascii=False)
                documents.append(RAGDocument(content=content, metadata=item_metadata))
        
        return documents
    
    def _process_csv_file(self, file_path: Path, metadata: Dict[str, Any]) -> List[RAGDocument]:
        """Process a CSV file."""
        df = pd.read_csv(file_path)
        documents = []
        
        for index, row in df.iterrows():
            # Convert row to string representation
            content = row.to_string()
            row_metadata = {**metadata, "row_index": index}
            documents.append(RAGDocument(content=content, metadata=row_metadata))
        
        return documents
    
    def _process_docx(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> RAGDocument:
        """Process a Word document and return a RAGDocument."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            if not text.strip():
                raise ValueError(f"Could not extract text from Word document {file_path}")
            
            # Create metadata if not provided
            if metadata is None:
                metadata = {}
            
            # Add source file information to metadata
            metadata.update({
                "source_file": os.path.basename(file_path),
                "file_type": "docx"
            })
            
            return RAGDocument(content=text.strip(), metadata=metadata)
            
        except Exception as e:
            raise ValueError(f"Error processing Word document {file_path}: {str(e)}")
    
    def _process_text(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> RAGDocument:
        """Process a text file and return a RAGDocument."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            if not text.strip():
                raise ValueError(f"Could not extract text from file {file_path}")
            
            # Create metadata if not provided
            if metadata is None:
                metadata = {}
            
            # Add source file information to metadata
            metadata.update({
                "source_file": os.path.basename(file_path),
                "file_type": "text"
            })
            
            return RAGDocument(content=text.strip(), metadata=metadata)
            
        except Exception as e:
            raise ValueError(f"Error processing text file {file_path}: {str(e)}")
    
    def _is_url(self, path: str) -> bool:
        """Check if a path is a URL."""
        try:
            result = urlparse(path)
            return all([result.scheme, result.netloc])
        except:
            return False
    
    def _process_url(self, url: str, metadata: Optional[Dict[str, Any]] = None) -> RAGDocument:
        """Process a URL and return a RAGDocument."""
        try:
            # Parse URL
            parsed_url = urlparse(url)
            domain = parsed_url.netloc
            
            # Handle Google Docs URLs
            if domain == 'docs.google.com':
                return self._process_google_doc(url, metadata)
            
            # Handle other URLs
            response = requests.get(url)
            response.raise_for_status()
            
            # Parse HTML
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Extract text content
            text = soup.get_text(separator='\n', strip=True)
            
            if not text.strip():
                raise ValueError(f"Could not extract text from URL {url}")
            
            # Create metadata if not provided
            if metadata is None:
                metadata = {}
            
            # Add source information to metadata
            metadata.update({
                "source_url": url,
                "domain": domain,
                "file_type": "web"
            })
            
            return RAGDocument(content=text.strip(), metadata=metadata)
            
        except Exception as e:
            raise ValueError(f"Error processing URL {url}: {str(e)}")
    
    def _process_google_doc(self, url: str, metadata: Optional[Dict[str, Any]] = None) -> RAGDocument:
        """Process a Google Doc URL and return a RAGDocument."""
        try:
            # Extract document ID from URL
            doc_id = url.split('/d/')[1].split('/')[0]
            
            # Construct export URL
            export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
            
            try:
                # Try to download content
                response = requests.get(export_url)
                response.raise_for_status()
                
                # Get text content
                text = response.text
                
                if not text.strip():
                    raise ValueError(f"Could not extract text from Google Doc {url}")
                
                # Create metadata if not provided
                if metadata is None:
                    metadata = {}
                
                # Add source information to metadata
                metadata.update({
                    "source_url": url,
                    "doc_id": doc_id,
                    "domain": "docs.google.com",
                    "file_type": "google_doc"
                })
                
                return RAGDocument(content=text.strip(), metadata=metadata)
                
            except requests.exceptions.HTTPError as e:
                if e.response.status_code == 401:
                    print("\nThis Google Doc requires authentication.")
                    print("Please make sure the document is publicly accessible or share it with the appropriate permissions.")
                    print("\nTo make the document accessible:")
                    print("1. Open the document in your browser")
                    print("2. Click 'Share' in the top right")
                    print("3. Click 'Change to anyone with the link'")
                    print("4. Set access to 'Viewer'")
                    print("5. Click 'Done'")
                    print("\nAfter updating the permissions, try running this command again.")
                    raise ValueError("Google Doc requires authentication. Please make it publicly accessible.")
                else:
                    raise e
            
        except Exception as e:
            raise ValueError(f"Error processing Google Doc {url}: {str(e)}")

    def _validate_file(self, file_path: str) -> bool:
        """Validate if a file can be processed.
        
        Args:
            file_path: Path to the file.
        
        Returns:
            True if the file is valid, False otherwise.
        """
        path = Path(file_path)
        if not path.exists():
            return False
        if path.suffix.lower() not in self.supported_formats:
            return False
        if path.stat().st_size > self.max_file_size:
            return False
        return True

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file.
        
        Returns:
            Extracted text from the PDF.
        """
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
            
        # If no text was extracted, try OCR
        if not text.strip():
            images = convert_from_path(file_path)
            for image in images:
                text += pytesseract.image_to_string(image) + "\n"
        
        return text

    def _extract_text(self, file_path: str) -> str:
        """Extract text from a file.
        
        Args:
            file_path: Path to the file.
        
        Returns:
            Extracted text from the file.
        """
        path = Path(file_path)
        if path.suffix.lower() == '.pdf':
            return self._extract_text_from_pdf(file_path)
        elif path.suffix.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks.
        
        Args:
            text: Text to split into chunks.
        
        Returns:
            List of text chunks.
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Adjust chunk end to not split words
            if end < len(text):
                while end > start and not text[end].isspace():
                    end -= 1
                if end == start:
                    end = start + self.chunk_size
            
            chunks.append(text[start:end].strip())
            start = end - self.chunk_overlap
        
        return [chunk for chunk in chunks if chunk]

    def _generate_document_id(self, content: str, metadata: Dict[str, Any]) -> str:
        """Generate a unique document ID.
        
        Args:
            content: Document content.
            metadata: Document metadata.
        
        Returns:
            Unique document ID.
        """
        unique_str = f"{content}{metadata}{uuid.uuid4()}"
        return hashlib.sha256(unique_str.encode()).hexdigest()[:32]

    async def process_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[str]:
        """Process a file and store it in the database.
        
        Args:
            file_path: Path to the file to process.
            metadata: Optional metadata for the document.
        
        Returns:
            List of document IDs for the processed chunks.
        """
        if not self._validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")

        metadata = metadata or {}
        metadata['source_file'] = file_path
        metadata['file_type'] = Path(file_path).suffix.lower()

        text = self._extract_text(file_path)
        chunks = self._chunk_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_index'] = i
            chunk_metadata['total_chunks'] = len(chunks)
            
            doc_id = self._generate_document_id(chunk, chunk_metadata)
            embedding = self.model.encode(chunk).tolist()
            
            doc = {
                'id': doc_id,
                'content': chunk,
                'metadata': chunk_metadata,
                'embedding': embedding
            }
            documents.append(doc)
        
        return await self.db.add_documents(documents)

    async def process_directory(self, dir_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[str]:
        """Process all supported files in a directory.
        
        Args:
            dir_path: Path to the directory to process.
            metadata: Optional metadata to apply to all documents.
        
        Returns:
            List of document IDs for all processed chunks.
        """
        doc_ids = []
        for root, _, files in os.walk(dir_path):
            for file in files:
                file_path = os.path.join(root, file)
                if self._validate_file(file_path):
                    try:
                        chunk_ids = await self.process_file(file_path, metadata)
                        doc_ids.extend(chunk_ids)
                    except Exception as e:
                        print(f"Error processing {file_path}: {e}")
        return doc_ids

    async def search(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Search for documents similar to the query.
        
        Args:
            query: Search query.
            limit: Maximum number of results to return.
        
        Returns:
            List of document dictionaries with similarity scores.
        """
        query_embedding = self.model.encode(query).tolist()
        return await self.db.search_documents(query_embedding, limit)

# Example usage
if __name__ == "__main__":
    # Initialize the document processor
    processor = DocumentProcessor()
    
    # Process a single file
    file_path = "sample.pdf"
    docs = processor.process_file(file_path)
    if docs:
        print(f"\nProcessed file: {file_path}")
        print("Content preview:")
        print("-" * 80)
        for doc in docs:
            print(doc.content[:200] + "...")
        print("-" * 80)
        print("\nMetadata:")
        print(docs[0].metadata)
    
    # Process a directory
    directory_path = "documents"
    docs = processor.process_directory(directory_path)
    print(f"\nProcessed {len(docs)} documents from {directory_path}") 